import logging
import re
from docx import Document

class WordEquationReplacer:
    def __init__(self, file_path, **kwargs):
        self.doc = Document(file_path)
        self.replacements = kwargs
        logging.debug(f"Инициализация WordEquationReplacer: {self.replacements}")

    import re

    def replace_text_in_paragraph(self, p):
        """Заменяет текст в переданном параграфе и логирует процесс."""
        for old, new in self.replacements.items():
            # Экранируем текст, чтобы избежать ошибок с регулярными выражениями
            sanitized_old = re.escape(old)

            # Используем re.sub для замены, просто указываем возможные разделители
            # Здесь мы обрабатываем возможные пробелы, запятые и точки
            pattern = re.compile(rf'\b{sanitized_old}\b')  # Слово должно быть целым, используем границы слова

            # Обработаем текст в параграфе
            if pattern.search(p.text):
                for run in p.runs:
                    if pattern.search(run.text):
                        old_text = run.text
                        run.text = pattern.sub(new, run.text)
                        logging.debug(f"Заменено в параграфе: '{old}' на '{new}' в тексте '{old_text}'")
            else:
                logging.debug(f"'{old}' не найдено в тексте: '{p.text}'")

    def replace_text_in_cell(self, cell):
        """Заменяет текст в переданной ячейке."""
        for paragraph in cell.paragraphs:
            self.replace_text_in_paragraph(paragraph)

    def process_document(self):
        """Обрабатывает документ, заменяя текст и обрабатывая таблицы."""
        logging.info("Обработка документа начата.")

        # Заменяем текст в параграфах
        for p in self.doc.paragraphs:
            logging.debug(f"Обрабатываем параграф: '{p.text}'")
            self.replace_text_in_paragraph(p)

        # Заменяем текст в таблицах
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    logging.debug(f"Обрабатываем ячейку: '{cell.text}'")
                    self.replace_text_in_cell(cell)

        logging.info("Обработка документа завершена.")

    def save_document(self, new_file_path):
        """Сохраняет документ в новый файл."""
        self.doc.save(new_file_path)
        logging.info(f"Документ сохранен как: {new_file_path}")
